import docx
from .base import DocumentParser, DocumentPage, TextBlock

def docx_table_to_markdown(table) -> str:
    """Convert python-docx table to markdown format."""
    rows_data = []
    for row in table.rows:
        row_cells = []
        for cell in row.cells:
            # Strip and replace newlines to keep markdown table row flat
            row_cells.append(cell.text.strip().replace("\n", " "))
        rows_data.append(row_cells)
    
    if not rows_data:
        return ""
        
    # De-duplicate rows for merged cells if needed (optional but good practice)
    # Simple markdown table generation:
    md = []
    header = rows_data[0]
    md.append("| " + " | ".join(header) + " |")
    md.append("| " + " | ".join(["---"] * len(header)) + " |")
    
    for row in rows_data[1:]:
        md.append("| " + " | ".join(row) + " |")
        
    return "\n".join(md)

class DocxParser(DocumentParser):
    def parse(self, path: str, include_pages: str | None = None) -> list[DocumentPage]:
        doc = docx.Document(path)
        
        text_blocks = []
        tables_out = []
        
        # Iterate through paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            font_size = 11.0  # Default font size
            is_bold = False
            
            # Check runs for font formatting
            sizes = []
            for run in paragraph.runs:
                if run.font.size:
                    sizes.append(run.font.size.pt)
                if run.bold:
                    is_bold = True
                    
            if sizes:
                font_size = sum(sizes) / len(sizes)
            else:
                # Fallback to paragraph style
                if paragraph.style and paragraph.style.font and paragraph.style.font.size:
                    font_size = paragraph.style.font.size.pt
            
            # Check style name for headings
            style_name = paragraph.style.name.lower() if paragraph.style else ""
            if "heading" in style_name:
                is_bold = True
                if "1" in style_name:
                    font_size = 20.0
                elif "2" in style_name:
                    font_size = 16.0
                elif "3" in style_name:
                    font_size = 14.0
                else:
                    font_size = 12.0
            
            text_blocks.append(TextBlock(
                text=text,
                font_size=font_size,
                is_bold=is_bold
            ))
            
        # Extract tables
        for table in doc.tables:
            md_table = docx_table_to_markdown(table)
            if md_table:
                tables_out.append(md_table)
                
        # Since DOCX doesn't have native physical pages, return everything as page 1
        page = DocumentPage(
            page_number=1,
            text_blocks=text_blocks,
            tables=tables_out,
            source_path=path
        )
        
        return [page]
